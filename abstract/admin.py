from django.contrib import admin
from import_export import resources, fields
from import_export.admin import ImportExportModelAdmin
from import_export.widgets import ForeignKeyWidget, ManyToManyWidget
from .models import *
from accounts.models import Users
from django.http import HttpResponse
from docx import Document
from django.core.mail import send_mail
from django.conf import settings
from .tasks import send_email_task
from django.utils.html import strip_tags

@admin.action(description="Export selected abstracts to DOCX")
def export_abstracts_to_docx(modeladmin, request, queryset):
    # Create a new Document
    doc = Document()
    doc.add_heading('Abstracts Export', level=1)

    for abstract in queryset:
        # Add the abstract title
        doc.add_heading(abstract.abstract_title, level=2)

        # Add submission ID and track
        doc.add_paragraph(f"Submission ID: {abstract.submission_id}")
        doc.add_paragraph(f"Track: {abstract.track.track_name}")

        # Add authors
        authors = abstract.authors.all()
        if authors:
            doc.add_paragraph("Authors:")
            for author in authors:
                doc.add_paragraph(
                    f"{author.author_name} ({author.affiliation}, {author.email})", style='List Bullet'
                )

        # Add presenters
        presenters = abstract.presentor.all()
        if presenters:
            doc.add_paragraph("Presenters:")
            for presenter in presenters:
                doc.add_paragraph(
                    f"{presenter.name} ({presenter.email})", style='List Bullet'
                )

        # Add other details
        doc.add_paragraph(f"Keywords: {', '.join(abstract.keywords.names())}")
        doc.add_paragraph(f"Status: {abstract.status}")
        
        # Add the abstract text itself
        doc.add_paragraph("Abstract:", style='Heading 3')
        doc.add_paragraph(strip_tags(abstract.abstract)) 

        # Divider between abstracts
        doc.add_paragraph("\n" + "-" * 40 + "\n")

    # Save the document to a response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=abstracts_export.docx'
    doc.save(response)

    return response



# Abstract Model Resource Class
class AbstractResource(resources.ModelResource):
    track = fields.Field(
        column_name='track',
        attribute='track',
        widget=ForeignKeyWidget(Track, 'track_name')
    )
    presentation_type = fields.Field(
        column_name='presentation_type',
        attribute='presentation_type',
        widget=ForeignKeyWidget(PresentationType, 'type_name')
    )
    user = fields.Field(
        column_name='Submitter',
        attribute='user',
        widget=ForeignKeyWidget(Users, 'email')
    )
    reviewers = fields.Field(
        column_name='reviewers',
        attribute='reviewers',
        widget=ManyToManyWidget(Reviewer, field='full_name')
    )

    # Include inlined fields for AuthorInformation and PresenterInformation
    author_names = fields.Field(column_name='author_names')
    author_emails = fields.Field(column_name='author_emails')
    author_affiliations = fields.Field(column_name='author_affiliations')
    presenter_names = fields.Field(column_name='presenter_names')
    presenter_emails = fields.Field(column_name='presenter_emails')

    def dehydrate_author_names(self, abstract):
        return "; ".join([author.author_name for author in abstract.authors.all()])

    def dehydrate_author_emails(self, abstract):
        return "; ".join([author.email for author in abstract.authors.all()])

    def dehydrate_author_affiliations(self, abstract):
        return "; ".join([author.affiliation for author in abstract.authors.all()])

    def dehydrate_presenter_names(self, abstract):
        return "; ".join([presenter.name for presenter in abstract.presentor.all()])

    def dehydrate_presenter_emails(self, abstract):
        return "; ".join([presenter.email for presenter in abstract.presentor.all()])

    class Meta:
        model = Abstract
        fields = (
            'id', 'abstract_title', 'abstract', 'submission_id', 'track', 'status',
            'presentation_type', 'user', 'keywords', 'attachment', 'date_created',
            'reviewers', 'author_names', 'author_emails', 'author_affiliations',
            'presenter_names', 'presenter_emails'
        )
        export_order = fields


# Reviewer Model Resource Class
class ReviewerResource(resources.ModelResource):
    email = fields.Field(
        column_name='email',
        attribute='email',
        widget=ForeignKeyWidget(Users, 'email')
    )
    expertise_area = fields.Field(
        column_name='expertise_area',
        attribute='expertise_area',
        widget=ForeignKeyWidget(Track, 'track_name')
    )

    class Meta:
        model = Reviewer
        fields = ('id', 'full_name', 'email', 'expertise_area')
        export_order = fields


# Assignment Model Resource Class
class AssignmentResource(resources.ModelResource):
    abstract = fields.Field(
        column_name='abstract',
        attribute='abstract',
        widget=ForeignKeyWidget(Abstract, 'abstract_title')
    )
    reviewer = fields.Field(
        column_name='reviewer',
        attribute='reviewer',
        widget=ForeignKeyWidget(Reviewer, 'full_name')
    )

    class Meta:
        model = Assignment
        fields = ('id', 'abstract', 'reviewer', 'status', 'assigned_at')
        export_order = fields


# Reviews Model Resource Class
class ReviewsResource(resources.ModelResource):
    abstract = fields.Field(
        column_name='abstract',
        attribute='abstract',
        widget=ForeignKeyWidget(Abstract, 'abstract_title')
    )
    user = fields.Field(
        column_name='user',
        attribute='user',
        widget=ForeignKeyWidget(Users, 'email')
    )
    reviewer = fields.Field(
        column_name='reviewer',
        attribute='reviewer',
        widget=ForeignKeyWidget(Reviewer, 'full_name')
    )
    presentation = fields.Field(
        column_name='presentation',
        attribute='presentation',
        widget=ForeignKeyWidget(PresentationType, 'type_name')
    )

    class Meta:
        model = Reviews
        fields = (
            'id', 'abstract', 'user', 'reviewer', 'comment', 'attachment', 'status',
            'presentation', 'date_created', 'title', 'content', 'relevance', 'quality', 
            'clarity', 'methods', 'structure', 'data_collection', 'result', 'conclusion', 'total'
        )
        export_order = fields


class EmailLogResource(resources.ModelResource):
    abstract = fields.Field(
        column_name='abstract',
        attribute='abstract',
        widget=ForeignKeyWidget(Abstract, 'abstract_title')
    )

    class Meta:
        model = EmailLog
        fields = ('recipient', 'abstract','subject', 'plain_message', 'sent_at')

# Admin Classes for Exporting

class AuthorInformationInline(admin.StackedInline):
    model = AuthorInformation
    extra = 1
    fields = ['author_name', 'email', 'affiliation']

class PresenterInformationInline(admin.StackedInline):
    model = PresenterInformation
    extra = 1
    fields = ['name', 'email']


class AssignmentInline(admin.TabularInline):
    model = Assignment
    extra = 0
    fields = ('abstract', 'status', 'assigned_at')
    readonly_fields = ('assigned_at',)


@admin.register(Abstract)
class AbstractAdmin(ImportExportModelAdmin):
    resource_class = AbstractResource
    list_display = ('abstract_title', 'submission_id', 'track', 'status', 'presentation_type')
    list_filter = ('status', 'track', 'presentation_type')
    search_fields = ('abstract_title', 'keywords', 'track__track_name')
    inlines = [AuthorInformationInline, PresenterInformationInline]
    ordering = ('-date_created',)
    actions = [export_abstracts_to_docx,'mark_as_accepted', 'mark_as_rejected', 'mark_as_reviewed']

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)

        # Send email ONLY if it's a new submission (not editing)
        if not change and obj.user:
            send_email_task.delay(
                email_type="abstract_submission",
                subject="Abstract Submission Confirmation",
                recipient=obj.user.email,
                template_name="emails/abstract_submission.html",
                context={"submission_id": obj.submission_id, "user_name": obj.user.first_name},
            )

    # 🔹 Bulk Action: Mark Abstracts as Accepted
    @admin.action(description="Mark selected abstracts as Accepted")
    def mark_as_accepted(self, request, queryset):
        self.bulk_update_status(queryset, 'Accepted', "Abstract Accepted")

    # 🔹 Bulk Action: Mark Abstracts as Rejected
    @admin.action(description="Mark selected abstracts as Rejected")
    def mark_as_rejected(self, request, queryset):
        self.bulk_update_status(queryset, 'Rejected', "Abstract Rejected")

    # 🔹 Bulk Action: Mark Abstracts as Reviewed
    @admin.action(description="Mark selected abstracts as Reviewed")
    def mark_as_reviewed(self, request, queryset):
        self.bulk_update_status(queryset, 'Reviewed', "Abstract Reviewed")

    # 🔥 Bulk Update Function
    def bulk_update_status(self, queryset, new_status, email_subject):
        for obj in queryset:
            obj.status = new_status
            obj.save()

            # Send email if user exists
            if obj.user:
                send_email_task.delay(
                    email_type=f"abstract_{new_status.lower()}",
                    subject=email_subject,
                    recipient=obj.user.email,
                    template_name="emails/abstract_status_update.html",
                    context={
                        "user_name": obj.user.first_name,
                        "abstract_title": obj.abstract_title,
                        "submission_id": obj.submission_id,
                        "new_status": new_status,
                    },
                )


@admin.register(Reviewer)
class ReviewerAdmin(ImportExportModelAdmin):
    resource_class = ReviewerResource
    list_display = ('full_name', 'email', 'expertise_area')
    search_fields = ('full_name', 'email__email')
    list_filter = ('expertise_area',)
    inlines = [AssignmentInline]

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)

        # Send email ONLY if it's a new reviewer registration (not editing)
        if not change and obj.email:
            send_email_task.delay(
                email_type="reviewer_registration",
                subject="Abstract Review Confirmation",
                recipient=obj.email,
                template_name="emails/reviewer_registration.html",
                context={"full_name": obj.full_name},
            )


@admin.register(Assignment)
class AssignmentAdmin(ImportExportModelAdmin):
    resource_class = AssignmentResource
    list_display = ('abstract', 'reviewer', 'status', 'assigned_at')
    list_filter = ('status', 'assigned_at')
    search_fields = ('abstract__abstract_title', 'reviewer__full_name')
    autocomplete_fields = ('abstract', 'reviewer')
    ordering = ('-assigned_at',)

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)

        # Send email ONLY if it's a new assignment (not editing)
        if not change and obj.user:
            send_email_task.delay(
                email_type="abstract_assignment",
                subject="New Abstract Assignment",
                recipient=obj.user.email,
                template_name="emails/abstract_assignment.html",
                context={
                    "full_name": obj.user.first_name,
                    "abstract_id": obj.abstract.submission_id,
                    "abstract_title": obj.abstract.abstract_title,
                },
            )


@admin.register(Reviews)
class ReviewsAdmin(ImportExportModelAdmin):
    resource_class = ReviewsResource
    list_display = ('abstract', 'user', 'reviewer', 'status', 'total', 'comment', 'attachment', 'presentation')
    list_filter = ('status', 'abstract', 'user')
    search_fields = ('abstract__abstract_title', 'user__username', 'comment')
    readonly_fields = ('total',)
    fieldsets = (
        ('Review Details', {
            'fields': ('abstract', 'user', 'reviewer', 'status', 'comment', 'attachment', 'presentation')
        }),
        ('Scoring', {
            'fields': (
                'title', 'content', 'relevance', 'quality', 'clarity',
                'methods', 'structure', 'data_collection', 'result',
                'conclusion', 'total'
            )
        }),
    )

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)

        # Send email ONLY if it's a new review submission (not editing)
        if not change and obj.user:
            send_email_task.delay(
                email_type="review_submission",
                subject="Review Submission Confirmation",
                recipient=obj.user.email,
                template_name="emails/review_submission.html",
                context={
                    "full_name": obj.user.first_name,
                    "abstract_id": obj.abstract.submission_id,
                    "abstract_title": obj.abstract.abstract_title,
                },
            )



@admin.register(EmailLog)
class EmailLogtAdmin(ImportExportModelAdmin):
    resource_class = EmailLogResource
    list_display = ('recipient', 'abstract','subject', 'plain_message', 'sent_at')
    list_filter = ('subject', 'sent_at')
    search_fields = ('abstract__abstract_title', 'recipient')
    autocomplete_fields = ('abstract', )
    ordering = ('-sent_at',)




# Register additional models
admin.site.register(Track)
admin.site.register(PresentationType)